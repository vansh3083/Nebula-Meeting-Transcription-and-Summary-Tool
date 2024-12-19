import os
from models.whisper_model import transcribe_audio
from models.summarizer_nlp import summarize_text  # Import summarization from summarizer_nlp.py
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph
from docx import Document
from docx.shared import Pt

def export_to_pdf(summary, keypoints, decisions, action_items, file_name):
    try:
        output_folder = "output"
        os.makedirs(output_folder, exist_ok=True) 
        pdf_path = os.path.join(output_folder, f"{file_name}.pdf")
        
        doc = SimpleDocTemplate(pdf_path, pagesize=letter)

        # Get default styles
        styles = getSampleStyleSheet()
        style_normal = styles['Normal']
  
        bold_style = ParagraphStyle(name='BoldHeading', fontName='Helvetica-Bold', fontSize=12)
        
        paragraphs = []

        paragraphs.append(Paragraph("Summary:", bold_style))
        paragraphs.append(Paragraph(summary, style_normal))
        
        paragraphs.append(Paragraph("\nKey Points:", bold_style))
        paragraphs.append(Paragraph(keypoints, style_normal))
        
        paragraphs.append(Paragraph("\nDecisions Made:", bold_style))
        paragraphs.append(Paragraph(decisions, style_normal))
        
        paragraphs.append(Paragraph("\nAction Items:", bold_style))
        paragraphs.append(Paragraph(action_items, style_normal))

        doc.build(paragraphs)

        print(f"Summary saved to PDF: {pdf_path}")
    except Exception as e:
        print(f"Error exporting to PDF: {e}")

# Function to export summary to Word document
def export_to_word(summary, keypoints, decisions, action_items, file_name):
    try:
        output_folder = "output"
        os.makedirs(output_folder, exist_ok=True)  # Ensure the output folder exists
        word_path = os.path.join(output_folder, f"{file_name}.docx")
        
        # Create a Word document with the summary, keypoints, decisions, and action items
        doc = Document()
        
        # Add headings with bold style
        doc.add_heading('Meeting Summary', 0)
        
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)
        
        doc.add_heading('Key Points', level=1)
        doc.add_paragraph(keypoints)
        
        doc.add_heading('Decisions Made', level=1)
        doc.add_paragraph(decisions)
        
        doc.add_heading('Action Items', level=1)
        doc.add_paragraph(action_items)
        
        # Save the document
        doc.save(word_path)
        print(f"Summary saved to Word: {word_path}")
    except Exception as e:
        print(f"Error exporting to Word: {e}")

# Main function to integrate transcription, summarization, and exporting
if __name__ == "__main__":
    
    audio_file = "data/audio/meeting2.mp3" 
    
    print("Starting transcription process...")
    transcript = transcribe_audio(audio_file)  # Get transcript from whisper_model.py

    if transcript:
        print("\nTranscription Output:")
        print(transcript)
        
        print("\nStarting summarization process...")
        summaries = summarize_text(transcript)  # Pass the transcript to summarizer
        
        if summaries:
            print("\nGenerated Summaries:")
            
            for model_name, result in summaries.items():
                print(f"\nSummary for {model_name}:")
                print(result["summary"])
                print("\nKey Points:")
                print(result["keypoints"])
                print("\nDecisions Made:")
                print(result["decisions"])
                print("\nAction Items:")
                print(result["action_items"])
                
                file_name = f"{os.path.splitext(os.path.basename(audio_file))[0]}_{model_name.replace('/', '_')}"
                
                # Export to PDF and Word
                export_to_pdf(result["summary"], result["keypoints"], result["decisions"], result["action_items"], file_name)
                export_to_word(result["summary"], result["keypoints"], result["decisions"], result["action_items"], file_name)
        else:
            print("Summarization failed. Please check the error messages above.")
    else:
        print("Transcription failed. Please check the error messages above.")
