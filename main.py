import os
from models.whisper_model import transcribe_audio
from models.summarizer import summarize_transcription, get_groq_chat_model
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from docx import Document

# Function to export summary to PDF
def export_to_pdf(summary, file_name):
    try:
        output_folder = "output"
        os.makedirs(output_folder, exist_ok=True)
        pdf_path = os.path.join(output_folder, f"{file_name}.pdf")

        doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        paragraphs = [Paragraph(summary, styles['Normal'])]

        doc.build(paragraphs)
        print(f"Summary saved to PDF: {pdf_path}")
    except Exception as e:
        print(f"Error exporting to PDF: {e}")


# Function to export summary to Word
def export_to_word(summary, file_name):
    try:
        output_folder = "output"
        os.makedirs(output_folder, exist_ok=True)
        word_path = os.path.join(output_folder, f"{file_name}.docx")

        doc = Document()
        doc.add_heading('Meeting Summary', 0)
        doc.add_paragraph(summary)
        doc.save(word_path)
        print(f"Summary saved to Word: {word_path}")
    except Exception as e:
        print(f"Error exporting to Word: {e}")


# Main function to integrate transcription, summarization, and exporting
if __name__ == "__main__":
    # Ensure the Groq API key is set in your environment
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        print("Error: GROQ_API_KEY environment variable is not set.")
        exit(1)  # Exit if the API key is not found

   
    audio_file = "data/video/meeting2.mp4"  # Default value if none is provided

    print("Starting transcription process...")
    transcript = transcribe_audio(audio_file)

    if transcript:
        print("\nTranscription Output:")
        print(transcript)

        print("\nInitializing LangChain Groq model...")
        model_name = input("Enter the Groq model name (e.g., 'llama3-8b-8192'): ").strip()
        chat_model = get_groq_chat_model(api_key, model_name=model_name)

        if chat_model:
            print("\nStarting summarization process...")
            summary = summarize_transcription(transcript, chat_model)

            if summary:
                print("\nGenerated Summary:")
                print(summary)

                # Export summary to PDF and Word
                file_name = os.path.splitext(os.path.basename(audio_file))[0] + "_groq"
                export_to_pdf(summary, file_name)
                export_to_word(summary, file_name)
            else:
                print("Summarization failed.")
        else:
            print("Failed to initialize LangChain Groq model.")
    else:
        print("Transcription failed.")
